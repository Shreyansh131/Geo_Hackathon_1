import os
import time
import streamlit as st
import speech_recognition as sr
import pyttsx3  # Note: pyttsx3 might not work in all environments, but kept for completeness
from streamlit_option_menu import option_menu
from pydantic import ValidationError
import json
import math
import numpy as np
import re  # Added for RegEx to infer Well ID from query
from langchain_community.vectorstores import Chroma


# --- Agent Imports (Pydantic models and Nodal Analysis functions are integrated for simplicity) ---
from typing import List, Dict, Any
from pydantic import BaseModel, Field

import math
import numpy as np

# --- Start of NodalAnalysis.py functions (integrated for single-file demonstration) ---
# --- Parallel & Logged Multimodal Loader (copy/paste replace your previous loader) ---
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import concurrent.futures
import hashlib
import json
import sys
from pathlib import Path
from typing import List, Dict

CACHE_PATH = Path("vectorDB") / ".index_cache.json"
SUPPORTED_EXT = (".pdf", ".docx", ".xlsx", ".png", ".jpg", ".jpeg", ".bmp", ".tiff")

def _file_mtime_map(root_dir: str) -> Dict[str, float]:
    m = {}
    for p in Path(root_dir).rglob("*"):
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXT:
            m[str(p)] = p.stat().st_mtime
    return m

def _log(msg: str, sidebar_container=None):
    """Log to terminal and optionally to Streamlit sidebar container."""
    ts = time.strftime("%H:%M:%S")
    line = f"[{ts}] {msg}"
    print(line)
    sys.stdout.flush()
    if sidebar_container:
        # maintain a session list for logs to persist across re-renders
        if 'scan_logs' not in st.session_state:
            st.session_state['scan_logs'] = []
        st.session_state['scan_logs'].append(line)
        # keep last 5000 chars to avoid UI bloat
        text = "\n".join(st.session_state['scan_logs'][-1000:])
        sidebar_container.text_area("Scanner logs", value=text, height=300)

# ---------- File parsers ----------
def parse_pdf(path: str) -> str:
    try:
        doc = fitz.open(path)
        pages = []
        for p in doc:
            pages.append(p.get_text())
        return "\n".join(pages)
    except Exception as e:
        return f"[PDF ERROR] {path}: {e}"

def parse_docx(path: str) -> str:
    try:
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"[DOCX ERROR] {path}: {e}"

def parse_xlsx(path: str) -> str:
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        out = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # convert None -> empty, then join columns with pipe for readability
                rows.append("| " + " | ".join("" if v is None else str(v) for v in row) + " |")
            if rows:
                out.append(f"### Sheet: {sheet}\n" + "\n".join(rows))
        return "\n\n".join(out)
    except Exception as e:
        return f"[XLSX ERROR] {path}: {e}"

def parse_image(path: str) -> str:
    # no OCR by default — simply tag the file so model knows there is an image
    return f"[IMAGE] {path} (OCR not applied)"

def _parse_file(path: str) -> Dict[str, str]:
    """Worker that returns a dict {path, text, ext}"""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        text = parse_pdf(path)
    elif ext == ".docx":
        text = parse_docx(path)
    elif ext == ".xlsx":
        text = parse_xlsx(path)
    else:
        text = parse_image(path)
    return {"path": path, "text": text, "ext": ext}

# ---------- Caching helpers ----------
def _load_cache() -> Dict:
    try:
        if CACHE_PATH.exists():
            return json.loads(CACHE_PATH.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}

def _save_cache(cache: Dict):
    try:
        CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
        CACHE_PATH.write_text(json.dumps(cache), encoding="utf-8")
    except Exception as e:
        print(f"[CACHE SAVE ERROR] {e}")

# ---------- Main loader ----------
def load_and_index_multimodal_data(root_dir: str, vectorstore: Chroma, max_workers: int = 4):
    """
    Parallel scanning + logging + caching loader.
    - root_dir: parent folder containing subfolders (wells)
    - vectorstore: session_state vectorstore or None (we will create new one)
    - max_workers: number of threads for parsing (adjust to your CPU)
    """

    sidebar_logs = st.sidebar.container()
    _log(f"Starting scan in '{root_dir}'", sidebar_container=sidebar_logs)

    if not os.path.isdir(root_dir):
        st.error(f"Directory not found: {root_dir}")
        return None

    # Discover files
    all_files = [str(p) for p in Path(root_dir).rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXT]
    total_files = len(all_files)
    if total_files == 0:
        st.error("No supported files found in folder.")
        return None

    _log(f"Discovered {total_files} files.", sidebar_container=sidebar_logs)
    st.info(f"Found {total_files} documents to index.")

    # Compute mtime map and check cache
    mtime_map = {p: Path(p).stat().st_mtime for p in all_files}
    cache = _load_cache()
    cached_mtimes = cache.get("mtimes", {})

    if cached_mtimes and cached_mtimes == mtime_map:
        # nothing changed -> fast path
        _log("No file changes since last index. Restoring existing vectorstore (fast path).", sidebar_container=sidebar_logs)
        if 'vectorstore' in st.session_state and st.session_state.vectorstore is not None:
            st.success("Index restored (no changes). Ready for questions.")
            st.session_state.indexing_done = True
            return st.session_state.vectorstore.as_retriever(search_kwargs={"k": 7})
        else:
            _log("Cache present but vectorstore not in session. Will rebuild.", sidebar_container=sidebar_logs)

    # Otherwise, we need to parse changed/new files (or reindex all)
    progress_bar = st.progress(0)
    status = st.empty()

    parsed_items: List[Dict[str, str]] = []

    # Workerize parsing with ThreadPoolExecutor (IO-bound)
    _log(f"Parsing files using {max_workers} workers...", sidebar_container=sidebar_logs)
    status.info("Parsing files...")

    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as exe:
        futures = {exe.submit(_parse_file, p): p for p in all_files}
        for i, fut in enumerate(concurrent.futures.as_completed(futures), start=1):
            p = futures[fut]
            try:
                res = fut.result()
                parsed_items.append(res)
                _log(f"[{i}/{total_files}] Parsed: {p}", sidebar_container=sidebar_logs)
            except Exception as e:
                parsed_items.append({"path": p, "text": f"[ERROR] {e}", "ext": Path(p).suffix.lower()})
                _log(f"[{i}/{total_files}] Failed: {p} -> {e}", sidebar_container=sidebar_logs)
            progress_bar.progress(i / total_files)
            status.write(f"Parsing file {i}/{total_files}: {os.path.basename(p)}")

    # Build Document objects and drop empty text
    from langchain_core.documents import Document  # local import to avoid top-level issues
    docs = []
    for item in parsed_items:
        text = item["text"] or ""
        if not text.strip():
            _log(f"Skipping empty output for {item['path']}", sidebar_container=sidebar_logs)
            continue
        meta = {"source": item["path"], "type": item["ext"]}
        docs.append(Document(page_content=text, metadata=meta))

    if not docs:
        st.error("No extractable textual content was parsed. Aborting.")
        return None

    status.info("Splitting into chunks...")
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    chunks = splitter.split_documents(docs)
    _log(f"Created {len(chunks)} chunks.", sidebar_container=sidebar_logs)
    st.info(f"Generated {len(chunks)} chunks. Running embedding smoke test...")

    # Embedding smoke test
    try:
        sample_text = chunks[0].page_content[:512]
        emb_out = st.session_state.embedding_model.embed_documents([sample_text])
        if not emb_out or not isinstance(emb_out, list) or len(emb_out[0]) == 0:
            raise RuntimeError("Embedding model returned invalid output.")
    except Exception as e:
        st.error(f"Embedding test failed: {e}")
        _log(f"Embedding test failed: {e}", sidebar_container=sidebar_logs)
        return None

    status.info("Embedding & indexing (this may take a while)...")
    _log("Starting Chroma indexing...", sidebar_container=sidebar_logs)

    try:
        new_vs = Chroma.from_documents(
            documents=chunks,
            embedding=st.session_state.embedding_model,
            persist_directory="vectorDB"
        )
        new_vs.persist()
        st.session_state.vectorstore = new_vs
        st.session_state.indexing_done = True
        _log("Chroma indexing complete.", sidebar_container=sidebar_logs)

        # Save cache mtimes so future runs can fast-path
        _save_cache({"mtimes": mtime_map})

        progress_bar.progress(1.0)
        status.success("Indexing complete. Ready for queries.")
        return new_vs.as_retriever(search_kwargs={"k": 7})

    except Exception as e:
        st.error(f"Indexing failed: {e}")
        _log(f"Indexing failed: {e}", sidebar_container=sidebar_logs)
        return None



def swamee_jain(Re, D, roughness):
    """Calculates the Darcy-Weisbach friction factor using the Swamee-Jain equation."""
    if Re <= 0:
        return 0.0
    return 0.25 / (math.log10((roughness / (3.7 * D)) + (5.74 / (Re ** 0.9)))) ** 2


def pump_interp(flow, pump_curve, key):
    """Interpolates pump head or flow from the provided pump curve data."""
    return np.interp(flow, pump_curve["flow"], pump_curve[key])


def perform_nodal_analysis(
        rho: float, mu: float, g: float = 9.81, roughness: float = 1e-5,
        reservoir_pressure: float = 230.0, wellhead_pressure: float = 10.0, PI: float = 5.0,
        esp_depth: float = 500.0, pump_curve: dict = None, well_trajectory: list = None,
        plot_results: bool = False
) -> dict:
    """Performs Nodal Analysis (IPR vs VLP) calculation and returns the operating point."""

    # Default pump curve if none is provided
    if pump_curve is None:
        pump_curve = {"flow": [0, 100, 200, 300, 400], "head": [600, 550, 450, 300, 100]}

    segments = []
    if not well_trajectory or len(well_trajectory) < 2:
        return {"status": "error", "message": "Invalid well trajectory provided."}

    for i in range(1, len(well_trajectory)):
        MD_seg = well_trajectory[i]["MD"] - well_trajectory[i - 1]["MD"]
        TVD_seg = well_trajectory[i]["TVD"] - well_trajectory[i - 1]["TVD"]
        D = well_trajectory[i]["ID"]
        L = MD_seg
        theta = math.atan2(TVD_seg, MD_seg) if MD_seg != 0 else math.pi / 2
        segments.append((L, D, theta))

    def vlp(flow_m3hr):
        q = flow_m3hr / 3600.0
        dp_total = 0.0
        depth_accum = 0.0

        for (L, D, theta) in segments:
            A = math.pi * D ** 2 / 4.0
            u = q / A
            Re = rho * abs(u) * D / mu
            f = swamee_jain(Re, D, roughness)

            dp_fric = f * (L / D) * (rho * u ** 2 / 2.0)
            dp_grav = rho * g * L * math.sin(theta)

            dp_total += dp_fric + dp_grav
            depth_accum += L * math.sin(theta)

        if depth_accum >= esp_depth:
            dp_total -= rho * g * pump_interp(flow_m3hr, pump_curve, "head")

        return wellhead_pressure + dp_total / 1e5

    def ipr(flow_m3hr):
        pbh = reservoir_pressure - flow_m3hr / PI
        return max(pbh, 0.0)

    flows = np.linspace(1, 400, 200)
    p_vlp = np.array([vlp(f) for f in flows])
    p_ipr = np.array([ipr(f) for f in flows])

    diff = np.abs(p_vlp - p_ipr)
    idx = np.argmin(diff)

    if diff[idx] < 3:
        sol_flow = flows[idx]
        sol_pbh = p_vlp[idx]
        sol_head = pump_interp(sol_flow, pump_curve, "head")

        results = {
            "status": "Solution Found",
            "flowrate_m3hr": float(f"{sol_flow:.2f}"),
            "pbh_bar": float(f"{sol_pbh:.2f}"),
            "pump_head_m": float(f"{sol_head:.1f}"),
        }
    else:
        results = {"status": "No solution found", "flowrate_m3hr": None}

    return results

# --- End of NodalAnalysis.py functions ---

# --- Integrated well_agents.py functions ---
class TrajectoryPoint(BaseModel):
    MD: float = Field(description="Measured Depth in meters.")
    TVD: float = Field(description="True Vertical Depth in meters.")
    ID: float = Field(description="Inner Diameter of casing/tubing in meters.")


class WellParameters(BaseModel):
    reservoir_pressure: float = Field(description="Reservoir pressure (bar).")
    wellhead_pressure: float = Field(description="Wellhead pressure (bar).")
    PI: float = Field(description="Productivity Index (m3/hr per bar).")
    esp_depth: float = Field(description="ESP intake depth (meters).")
    well_trajectory: List[TrajectoryPoint] = Field(description="List of casing/tubing segments.")

    rho: float = 1000.0
    mu: float = 1e-3
    roughness: float = 1e-5
    pump_curve: dict = Field(default={"flow": [0, 100, 200, 300, 400], "head": [600, 550, 450, 300, 100]})


def validate_well_parameters(params: WellParameters) -> dict:
    validation_report = {"errors": [], "warnings": [], "validated_params": params.dict()}

    for i, point in enumerate(params.well_trajectory):
        if point.MD < point.TVD - 1.0:
            validation_report["errors"].append(
                f"Critical Error: MD ({point.MD:.1f}m) is significantly less than TVD ({point.TVD:.1f}m) at segment {i}. Check for unit mix-up or error."
            )

    if params.well_trajectory:
        final_tvd = params.well_trajectory[-1].TVD
        if final_tvd < params.esp_depth:
            validation_report["errors"].append(
                f"Critical Error: Final TVD ({final_tvd:.1f}m) is shallower than ESP depth ({params.esp_depth:.1f}m). Pump location is invalid."
            )

    if params.PI < 0.1 or params.PI > 50.0:
        validation_report["warnings"].append(
            f"Warning: PI ({params.PI:.1f}) is outside typical range (0.1 - 50)."
        )

    return validation_report


def run_nodal_analysis_tool(validated_params: WellParameters) -> dict:
    try:
        results = perform_nodal_analysis(**validated_params.dict(), plot_results=False)
        return {"status": "success", "results": results}
    except Exception as e:
        return {"status": "error", "message": f"Nodal Analysis failed during execution: {str(e)}"}


# --- End of integrated well_agents.py functions ---

# --- Existing LangChain/Ollama Imports ---
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.llms import Ollama

# Core prompt template
from langchain_core.prompts import PromptTemplate

# Memory (from langchain-classic)
from langchain_classic.memory import ConversationBufferMemory

# Chains (from langchain-classic)
from langchain_classic.chains import RetrievalQA

# Text splitter (from langchain-text-splitters)
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Community integrations
from langchain_community.chat_message_histories import StreamlitChatMessageHistory
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings.ollama import OllamaEmbeddings
from langchain_community.llms import Ollama
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from unstructured.partition.auto import partition
import speech_recognition as sr

try:
    engine = pyttsx3.init('sapi5')
    voices = engine.getProperty('voices')
    engine.setProperty('voice', voices[1].id)


    def speak(audio):
        engine.say(audio)
        engine.runAndWait()


    def SpeakNow(command):
        voice = pyttsx3.init()
        voice.say(command)
        voice.runAndWait()
except Exception:
    def speak(audio):
        pass


    def SpeakNow(command):
        pass
audio = sr.Recognizer()

# --- LLM and Model Configuration ---
TEXT_MODEL = "llama3"
VISION_MODEL = "llava-phi3"
OLLAMA_BASE_URL = "http://localhost:11434"

# Session state setup
if 'template' not in st.session_state:
    st.session_state.template = """You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.

Context: {context}
History: {history}

User: {question}
Chatbot:"""

if 'prompt' not in st.session_state:
    st.session_state.prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template=st.session_state.template,
    )

if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question",
        chat_memory=StreamlitChatMessageHistory()
    )

if 'text_llm' not in st.session_state:
    st.session_state.text_llm = Ollama(
        base_url=OLLAMA_BASE_URL,
        model=TEXT_MODEL,
        verbose=True
    )

if 'vision_llm' not in st.session_state:
    st.session_state.vision_llm = Ollama(
        base_url=OLLAMA_BASE_URL,
        model=VISION_MODEL,
        verbose=True
    )

if 'embedding_model' not in st.session_state:
    st.session_state.embedding_model = OllamaEmbeddings(base_url=OLLAMA_BASE_URL, model=TEXT_MODEL)

if 'vectorstore' not in st.session_state:
    st.session_state.vectorstore = Chroma(
        persist_directory='vectorDB',
        embedding_function=st.session_state.embedding_model
    )

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# UI
st.header("Agentic Nodal Analysis Assistant 🤖")


def multipage_menu(caption):
    st.sidebar.title(caption)
    st.sidebar.subheader('Navigation')


multipage_menu("Agentic AI Assistant")

selected = option_menu(
    menu_title="Main Menu",
    options=["Home", "History", "Voice Text"],
    icons=["house", "book", "mic"],
    menu_icon="cast",
    default_index=0,
    orientation="horizontal",
)

if selected == "Voice Text":
    try:
        with sr.Microphone() as source2:
            with st.spinner('Listening...'):
                time.sleep(1)
            audio.adjust_for_ambient_noise(source2, duration=1)
            st.write("Speak now please...")
            audio2 = audio.listen(source2)
            text = audio.recognize_google(audio2).lower()
            say = "Did you say " + text
            SpeakNow(say)
            st.chat_input(text, key="voice_output_input")
    except sr.WaitTimeoutError:
        st.warning("No speech detected. Please try again.")
    except Exception as e:
        st.error(f"Speech recognition error: {e}")

if selected == "History":
    st.subheader("Chat History")
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["message"])
    st.subheader("Indexed Files")
    st.info("The index stores data from the folder you specified on the Home tab.")


# --- Multimodal Indexing Function (UPDATED FOR MULTI-WELL) ---
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import io

def load_and_index_multimodal_data(well_dir: str, vectorstore: Chroma):
    """Loads PDF, DOCX, XLSX, images using lightweight Windows-safe parsers."""

    st.info("Scanning well folder...")

    supported_ext = ('.pdf', '.docx', '.xlsx', '.png', '.jpg', '.jpeg')
    file_list = []

    for root, _, files in os.walk(well_dir):
        for file in files:
            if file.lower().endswith(supported_ext):
                file_list.append(os.path.join(root, file))

    if not file_list:
        st.error("No supported documents found.")
        return None

    st.success(f"Found {len(file_list)} documents.")

    docs = []

    # --------------------------------------------------------------------------------
    # PDF PARSING (FAST — PyMuPDF)
    # --------------------------------------------------------------------------------
    def parse_pdf(path):
        try:
            pdf = fitz.open(path)
            text = ""
            for page in pdf:
                text += page.get_text()
            return text
        except Exception as e:
            return f"[PDF ERROR] Could not parse {path}: {e}"

    # --------------------------------------------------------------------------------
    # DOCX PARSING (python-docx)
    # --------------------------------------------------------------------------------
    def parse_docx(path):
        try:
            doc = DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            return f"[DOCX ERROR] Could not parse {path}: {e}"

    # --------------------------------------------------------------------------------
    # XLSX PARSING (openpyxl → markdown-style tables)
    # --------------------------------------------------------------------------------
    def parse_xlsx(path):
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                table_md = f"\n### Sheet: {sheet}\n\n"
                for row in ws.iter_rows(values_only=True):
                    row = [" " if v is None else str(v) for v in row]
                    table_md += "| " + " | ".join(row) + " |\n"
                text += table_md
            return text
        except Exception as e:
            return f"[XLSX ERROR] Could not parse {path}: {e}"

    # --------------------------------------------------------------------------------
    # IMAGE PARSING (no OCR by default — optional)
    # --------------------------------------------------------------------------------
    def parse_image(path):
        return f"[IMAGE FOUND] {path} — No OCR performed. Provide description if needed."

    # --------------------------------------------------------------------------------
    # PROCESS FILES
    # --------------------------------------------------------------------------------
    all_docs = []
    for path in file_list:
        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            content = parse_pdf(path)
        elif ext == ".docx":
            content = parse_docx(path)
        elif ext == ".xlsx":
            content = parse_xlsx(path)
        else:
            content = parse_image(path)

        all_docs.append(
            Document(
                page_content=content,
                metadata={"source": path, "type": ext}
            )
        )

    # --------------------------------------------------------------------------------
    # CHUNKING
    # --------------------------------------------------------------------------------
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    chunks = splitter.split_documents(all_docs)

    # --------------------------------------------------------------------------------
    # VECTORSTORE UPDATE
    # --------------------------------------------------------------------------------
    st.info("Embedding & indexing...")

    new_vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=st.session_state.embedding_model,
        persist_directory='vectorDB'
    )

    new_vectorstore.persist()
    st.session_state.vectorstore = new_vectorstore

    st.success(f"Indexed {len(chunks)} chunks successfully.")

    return new_vectorstore.as_retriever(search_kwargs={"k": 7})

# --- LLM Data Extraction Helper (VISION AGENT'S TASK) ---
def extract_structured_params(query: str, retriever) -> str:
    """
    Uses the VISION LLM to extract parameters into a JSON string,
    applying a metadata filter for the target well ID.
    """

    # 1. Infer Well ID from the query
    # Look for patterns like "Well 1", "Well 12", "Well-A", etc.
    # This regex is broad to capture various naming conventions (e.g., "well3", "well 10", "well-abc")
    match = re.search(r'(well\s*[\w\d]+)', query, re.IGNORECASE)
    well_id_raw = match.group(0).replace('well', '').strip() if match else None

    if not well_id_raw:
        st.warning("Please specify a target well (e.g., 'Well 3') in your Nodal Analysis query.")
        return '{"error": "Well ID missing"}'

    # Normalization: Capitalize the first letter for consistent metadata filtering (e.g., 'well 1' -> 'Well 1')
    # This logic assumes the folder name matches this normalized format (e.g., 'Well 1' vs 'well 1')
    well_id = well_id_raw[0].upper() + well_id_raw[1:] if len(well_id_raw) > 0 else well_id_raw

    st.info(f"Applying metadata filter for target well: {well_id}")

    # 2. Build Metadata Filter (CRITICAL FOR PERFORMANCE)
    metadata_filter = {"well_id": well_id}

    # 3. Retrieval with Filter
    search_query = query + " PI, reservoir pressure, wellhead pressure, TVD, tubing ID"

    # Chroma retrieval syntax for filtering by metadata
    retrieved_docs = st.session_state.vectorstore.as_retriever(
        search_kwargs={"k": 7, "where": metadata_filter}
    ).invoke(search_query)

    if not retrieved_docs:
        return '{"error": "No relevant documents found for the specified well and parameters."}'

    context = "\n---\n".join([doc.page_content for doc in retrieved_docs])

    # 4. VISION LLM Call for Structured Extraction
    extraction_prompt = f"""
    You are an expert petroleum data extraction agent. Extract the following well parameters from the context below, 
    which pertains ONLY to Well ID: {well_id}. Format the output STRICTLY as a JSON object that matches the following schema.
    If a required value is missing or ambiguous, use the placeholder 0.0 and generate a warning in the validation stage later.

    Pydantic Schema (for output structure):
    {WellParameters.schema_json(indent=2)}

    Context (Retrieved Documents):
    {context}

    Extract parameters for Nodal Analysis, outputting ONLY the JSON object:
    """

    response = st.session_state.vision_llm.invoke(extraction_prompt)

    try:
        json_str = response.strip()
        if '```json' in json_str:
            json_str = json_str.split("```json")[1].split("```")[0]
        return json_str
    except IndexError:
        return response

    # --- Home Tab Logic (The Router Agent) ---


if selected == "Home":
    st.title(f"You have selected {selected}")
    col1, col2 = st.columns(2)
    with col1:
        # --- FIX: Replaced problematic st.image with icon/markdown ---
        st.markdown("<p style='font-size: 64px;'>🛢️</p>", unsafe_allow_html=True)
        st.caption("Agentic Workflow Diagram")
    with col2:
        st.title("Well Performance Optimizer")

    # Change the input to accept the PARENT folder name
    parent_folder_name = st.text_input(
        "Enter the name of the Parent Data Folder (e.g., 'Wells_Archive'):",
        key="parent_folder_input",
        value="Wells_Archive"  # Default value updated for multi-well structure
    )

    if st.button("Index All Well Data"):
        # Create dummy folders and vectorDB if they don't exist
        os.makedirs(parent_folder_name, exist_ok=True)
        os.makedirs('vectorDB', exist_ok=True)

        # Correctly uses user input to find the parent data path
        parent_path = f'./Training data-shared with participants'

        if os.path.isdir(parent_path):
            st.session_state.retriever = load_and_index_multimodal_data(parent_path, st.session_state.vectorstore)
            st.session_state.indexing_done = True
        else:
            st.error(
                f"Directory not found: {parent_path}. Please create a directory named '{parent_folder_name}' and ensure it contains sub-folders for each well (e.g., 'Well 1').")

    if st.session_state.get('indexing_done', False):
        st.info(f"Data indexed. Q&A uses **{TEXT_MODEL}** (Fast). Calculation uses **{VISION_MODEL}** (Accurate).")

        # Retriever here is now an unfiltered retriever for initial Q&A
        st.session_state.retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": 7})

        if 'qa_chain' not in st.session_state:
            st.session_state.qa_chain = RetrievalQA.from_chain_type(
                llm=st.session_state.text_llm,
                chain_type='stuff',
                retriever=st.session_state.retriever,  # This will search all wells unless filtered
                verbose=False,
                chain_type_kwargs={
                    "verbose": False,
                    "prompt": st.session_state.prompt,
                    "memory": st.session_state.memory,
                }
            )

        if user_input := st.chat_input(
                "Ask a general question, or request a Nodal Analysis calculation (e.g., 'Estimate production rate for Well 3'):",
                key="user_input"):
            user_message = {"role": "user", "message": user_input}
            st.session_state.chat_history.append(user_message)
            with st.chat_message("user"):
                st.markdown(user_input)

            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                full_response = ""

                # --- ROUTER AGENT: DECISION LOGIC ---
                if "production rate" in user_input.lower() or "nodal analysis" in user_input.lower() or "flow rate" in user_input.lower():
                    # --- PATH A: CALCULATION / VISION AGENT (Requires Well ID filtering) ---

                    # 1. Check for Well ID in query
                    match = re.search(r'(well\s*[\w\d]+)', user_input, re.IGNORECASE)
                    if not match:
                        final_response = "**Error:** Please specify which well you want to analyze (e.g., 'Well 5')."
                    else:
                        full_response += f"Query detected for Nodal Analysis (Target: {match.group(0)}). Using VISION AGENT...\n"
                        message_placeholder.markdown(full_response + "▌")

                        with st.status(
                                "Running Agent Workflow: Data Retrieval -> Validation -> Calculation (Filtered by Well ID)"):

                            st.write("1. Retrieving data and structuring parameters using **Vision Agent**...")
                            # extract_structured_params handles the metadata filtering based on the inferred well ID
                            extracted_json_str = extract_structured_params(user_input, st.session_state.retriever)

                            if extracted_json_str == '{"error": "Well ID missing"}':
                                final_response = "**Error:** Please specify the target well in your query (e.g., 'Well 3')."
                            elif extracted_json_str == '{"error": "No relevant documents found for the specified well and parameters."}':
                                final_response = "**Error:** Could not find any relevant technical data for that well. Check indexing status."
                            else:
                                try:
                                    extracted_data = json.loads(extracted_json_str)

                                    st.write("2. Pydantic validation and type casting...")
                                    mock_params = WellParameters(**extracted_data)

                                    st.write("3. Domain sanity checking (MD vs TVD, PI range)...")
                                    validation_report = validate_well_parameters(mock_params)

                                    if validation_report['errors']:
                                        final_response = f"**Calculation Halted (Validation Failed):** Critical errors found in extracted data: {validation_report['errors']}. Review the well report data."
                                    else:
                                        st.write("4. Executing Nodal Analysis script...")
                                        calc_results = run_nodal_analysis_tool(mock_params)

                                        if calc_results['status'] == 'success':
                                            Q = calc_results['results'].get('flowrate_m3hr')
                                            BHP = calc_results['results'].get('pbh_bar')
                                            final_response = f"**Nodal Analysis Complete for {match.group(0)}!**\n\n- **Estimated Flowrate (Q):** **{Q:.2f} m³/hr**\n- **Bottomhole Pressure (BHP):** **{BHP:.2f} bar**\n\nWarnings: {validation_report['warnings']}"
                                        else:
                                            final_response = f"Calculation failed: {calc_results['message']}"

                                except (json.JSONDecodeError, ValidationError) as e:
                                    final_response = f"**Agent Error:** Failed to parse JSON or Pydantic validation failed during extraction. Error: {e}"

                        response_text = final_response

                else:
                    # --- PATH B: STANDARD Q&A / TEXT AGENT ---
                    with st.spinner(f"Thinking answer using the fast {TEXT_MODEL} agent..."):
                        response = st.session_state.qa_chain(user_input)
                    response_text = response['result']

                # Final rendering
                for chunk in response_text.split():
                    full_response += chunk + " "
                    time.sleep(0.02)
                    message_placeholder.markdown(full_response + "▌")
                message_placeholder.markdown(full_response)

                chatbot_message = {"role": "assistant", "message": response_text}
                st.session_state.chat_history.append(chatbot_message)

    else:
        st.write("Please enter the parent folder name and click 'Index All Well Data' to start the agent.")